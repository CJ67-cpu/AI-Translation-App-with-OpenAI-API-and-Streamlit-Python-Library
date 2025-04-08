
import streamlit as st
from openai import OpenAI
from docx import Document
from langdetect import detect, DetectorFactory
from io import BytesIO
import re
import tiktoken
import pandas as pd
from prompt_modules_short import (
    gender_module, genre_style_module, consistency_module,
    dialogue_module, idiom_module, formatting_module
)

DetectorFactory.seed = 0

if "OPENAI_API_KEY" not in st.secrets:
    st.error("🔐 OPENAI_API_KEY is not set. Please add it in your Streamlit Cloud app settings under 'Secrets'.")
    st.stop()

client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

st.set_page_config(page_title="Book Translator")

st.title("📚 Book Translator")
st.markdown("Upload a plain text (.txt) or Word (.docx) file in any language. This app will detect the language and translate it into English using GPT-4 or GPT-3.5 Turbo.")

model_choice = st.selectbox("Choose translation quality:", ["GPT-3.5 Turbo", "GPT-4 Turbo"])
model = "gpt-3.5-turbo" if model_choice == "GPT-3.5 Turbo" else "gpt-4"
max_model_tokens = 4096 if model == "gpt-3.5-turbo" else 16385

custom_instruction = st.text_input(
    "Optional: Add a style or tone instruction (e.g. 'translate like a gothic novel')",
    placeholder="Leave blank for default translation"
)

uploaded_file = st.file_uploader("Upload a text or Word file", type=["txt", "docx"])

def read_uploaded_file(uploaded_file):
    if uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        st.error("Unsupported file format.")
        return ""

def build_prompt(chunk, detected_lang, custom_instruction, gender_sensitive):
    injected_modules = [
        consistency_module,
        dialogue_module,
        idiom_module,
        formatting_module
    ]

    if gender_sensitive:
        injected_modules.insert(0, gender_module)

    if custom_instruction.strip():
        injected_modules.append(genre_style_module(custom_instruction.strip()))

    translation_instructions = "\n\n".join(injected_modules)

    prompt = f"""You are a literary translator. Translate this {detected_lang} text into English.

{translation_instructions}

Text:
{chunk}

English Translation:"""
    return prompt

def subchunk_if_needed(chunk, prompt_tokens, max_total_tokens, enc):
    paragraphs = re.split(r'\n\s*\n', chunk)
    subchunks = []
    current = ""
    current_tokens = 0

    for para in paragraphs:
        para_tokens = len(enc.encode(para))
        if current_tokens + para_tokens <= max_total_tokens - prompt_tokens:
            current += para + "\n\n"
            current_tokens += para_tokens
        else:
            if current:
                subchunks.append(current.strip())
            current = para + "\n\n"
            current_tokens = para_tokens
    if current:
        subchunks.append(current.strip())

    return subchunks

if uploaded_file is not None:
    raw_text = read_uploaded_file(uploaded_file)
    st.success("File uploaded successfully.")

    detected_lang = detect(raw_text)
    st.write(f"🌍 Detected language: **{detected_lang.upper()}**")

    with st.expander("View original text"):
        st.text_area("Original text (first 500 characters):", raw_text[:500], height=200)

    word_count = len(raw_text.split())
    estimated_tokens = word_count * 2

    if model == "gpt-3.5-turbo":
        cost = (estimated_tokens / 1000) * (0.0015 + 0.002)
    else:
        cost = (estimated_tokens / 1000) * (0.01 + 0.03)

    st.info(f"Estimated translation cost using {model_choice}: **${cost:.2f}** for ~{word_count} words")

    paragraphs = re.split(r'\n\s*\n', raw_text)
    initial_chunks = []
    temp = []
    word_limit = 300
    for para in paragraphs:
        temp.append(para)
        if len(" ".join(temp).split()) >= word_limit:
            initial_chunks.append("\n\n".join(temp))
            temp = []
    if temp:
        initial_chunks.append("\n\n".join(temp))

    if st.button("Translate Text"):
        st.info("Translating... This may take a while for large files.")
        translated_chunks = []
        token_stats = []
        progress = st.progress(0)
        enc = tiktoken.encoding_for_model(model)
        gender_sensitive = detected_lang in ["es", "it", "pt", "fr"]

        total = len(initial_chunks)
        count = 0

        for chunk in initial_chunks:
            prompt = build_prompt(chunk, detected_lang, custom_instruction, gender_sensitive)
            token_count = len(enc.encode(prompt))
            char_count = len(prompt)

            if token_count > max_model_tokens:
                subchunks = subchunk_if_needed(chunk, prompt_tokens=1000, max_total_tokens=max_model_tokens, enc=enc)
            else:
                subchunks = [chunk]

            for sub in subchunks:
                prompt = build_prompt(sub, detected_lang, custom_instruction, gender_sensitive)
                token_count = len(enc.encode(prompt))
                char_count = len(prompt)

                token_stats.append({
                    "Chunk": count + 1,
                    "Characters": char_count,
                    "Tokens Used": token_count,
                    "Max Allowed": max_model_tokens
                })

                try:
                    response = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant who translates books."},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    translation = response.choices[0].message.content
                    translated_chunks.append(translation)
                except Exception as e:
                    st.error(f"Error translating part {count+1}: {e}")
                    translated_chunks.append("[Translation failed for this part]")
                count += 1
                progress.progress(count / total)

        full_translation = "\n\n".join(translated_chunks)

        st.subheader("✅ Translated Text")
        st.text_area("Translation (preview):", full_translation[:1500], height=300)

        doc = Document()
        doc.add_heading("Translated Document", 0)
        for paragraph in full_translation.split("\n\n"):
            doc.add_paragraph(paragraph)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Download Translated .docx",
            data=buffer,
            file_name="translated_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Show token usage stats
        st.subheader("📊 Token Usage Per Chunk")
        df = pd.DataFrame(token_stats)
        st.dataframe(df)
